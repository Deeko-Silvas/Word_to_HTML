from docx import Document
import os

element_dict = {"Normal":"p", "Heading 1":"h1", "Heading 2":"h2", "Heading 3":"h3", "Heading 4":"h4", "List Paragraph":"li"}
# menu_file = "C:/Users/phill/OneDrive/Documents/Web Development/Cheatsheets/Front End Development Notes/menu.html"
menu_file = "Front End Development Notes/menu.html"
file_structure_list = []
sub_folder_index = 0

# doc = Document("C:/Users/phill/OneDrive/Documents/Web Development/Cheatsheets/Front End Development Notes.docx")
doc = Document("Front End Development Notes.docx")
rels = doc.part.rels

for index, para in enumerate(doc.paragraphs):
    style = para.style
    text = para.text
    if not text.isspace() and para.text != "":
        if index == 0 and style.name == "Heading 1":
            # dirName = f"C:/Users/phill/OneDrive/Documents/Web Development/Cheatsheets/{text}"
            dirName = text
            file_structure_list.append([text])
            if not os.path.exists(dirName):
                os.mkdir(dirName)
            pass
        elif style.name == "Heading 1":
            subDirName = f"{dirName}/{text}"
            file_structure_list[0].append([text])
            sub_folder_index += 1
            if not os.path.exists(subDirName):
                os.mkdir(subDirName)
        elif style.name =="Heading 2":
            file_name_camel = text.replace(" ", "")
            fileName = f"{subDirName}/{file_name_camel}.html"
            file_structure_list[0][sub_folder_index].append(text)
            try:
                file.close()
            except NameError:
                None
            file = open(fileName, "w")
            e = element_dict[style.name]
            file.write(f"<{e}>{text}</{e}>")
        elif "toc" not in style.name[:3] and text != "":
            e = element_dict[style.name]
            text = text.replace("<", "&lt;")
            text = text.replace(">", "&gt;")
            text = text.replace("”", "&ldquo;")
            file.write(f"<{e}>{text}</{e}>")

def create_menu():
    file = open(menu_file, "w")
    file.write('<ul class="sidebar-menu">')
    for n, nested1 in enumerate(file_structure_list[0]):
        if n > 0:
            file.write(f'<li class="treeview"><a href="#"><span>{file_structure_list[0][n][0]}</span><i class="fa fa-angle-left float-right"></i></a>')
            file.write('<ul class="treeview-menu">')
            for i, nested2 in enumerate(file_structure_list[0][n]):
                if i > 0:
                    nested_camel = nested2.replace(" ", "")
                    file.write(f'<li id="{nested_camel}"><a href="#"><i class="fa fa-circle-o"></i>{nested2}</a></li>')
            file.write('</ul>')
            file.write('</li>')
    file.write('</ul>')

create_menu()